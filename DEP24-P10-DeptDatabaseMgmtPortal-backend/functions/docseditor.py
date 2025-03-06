from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt,RGBColor

achievement_records = [
    {'faculty': 'Dr. A', 'title': 'Best faculty award', 'type': 'National'},
    {'faculty': 'Dr. B', 'title': 'Best faculty award', 'type': 'National'},
    {'faculty': 'Dr. C', 'title': 'Best faculty award', 'type': 'National'},
    {'faculty': 'Dr. D', 'title': 'Best faculty award', 'type': 'National'},
]

student_achievement_records = [
    {'student': 'A', 'title': 'Best faculty award', 'type': 'National'},
    {'student': 'B', 'title': 'Best faculty award', 'type': 'National'},
    {'student': 'C', 'title': 'Best faculty award', 'type': 'National'},
    {'student': 'D', 'title': 'Best faculty award', 'type': 'National'},
]

event_records = [
    {'speakers': 'Dr. A', 'title': 'Best faculty award', 'date': '2020-01-01', 'venue': 'Delhi', 'number of participants': '100', 'description': 'National'},
    {'speakers': 'Dr. B', 'title': 'Best faculty award', 'date': '2020-01-01', 'venue': 'Delhi', 'number of participants': '100', 'description': 'National'},
    {'speakers': 'Dr. C', 'title': 'Best faculty award', 'date': '2020-01-01', 'venue': 'Delhi', 'number of participants': '100', 'description': 'National'},
    {'speakers': 'Dr. D', 'title': 'Best faculty award', 'date': '2020-01-01', 'venue': 'Delhi', 'number of participants': '100', 'description': 'National'},
]

visit_records = [
    {'user': 'Dr. A', 'title': 'Best faculty award', 'from_date': '2020-01-01', 'venue': 'Delhi', 'description': 'National'},
    {'user': 'Dr. B', 'title': 'Best faculty award', 'from_date': '2020-01-01', 'venue': 'Delhi', 'description': 'National'},
    {'user': 'Dr. C', 'title': 'Best faculty award', 'from_date': '2020-01-01', 'venue': 'Delhi', 'description': 'National'},
    {'user': 'Dr. D', 'title': 'Best faculty award', 'from_date': '2020-01-01', 'venue': 'Delhi', 'description': 'National'},
]

Department = {
    'name': 'Computer Science and Engineering',
    'Hod': {
    'username':'rajesh'
    }
}



def get_admin_staff_count(staff):
    return 10

def get_tech_staff_count(staff):
    return 10

def BTech_students_count(students):
    count = 0
    for i in students:
        if i['programme'] == 'BTech':
            count += 1
    return count

def MTech_students_count(students):
    count = 0
    for i in students:
        if i['programme'] == 'MTech':
            count += 1
    return count

def MS_students_count(students):
    count = 0
    for i in students:
        if i['programme'] == 'MS':
            count += 1
    return count

def PhD_students_count(students):
    count = 0
    for i in students:
        if i['programme'] == 'PhD':
            count += 1
    return count

def get_faculty_count(faculties):
    return len(faculties)

def get_publications_count(publications_records):
    return 10

def get_ug_labs_counts(labs):
    return 10

def get_pg_labs_counts(labs):
    return 10

def get_research_labs_counts(labs):
    return 10



# filename, Department, faculties, students, staff, research_labs, achievement_records,student_achievement_records,event_records,visit_records, publications_records, year

faculties = []
students = []
staff = []
publications_records = []
research_labs = None
year = 2019

document = Document('editfile.docx')
section = document.sections[0]
header = section.header
header.paragraphs[0].add_run(f'Annual Report {year}-{(year+1)%100}\n')

# Adding Title
document.add_heading(f'DEPARTMENT OF {Department["name"]}'.upper(), 0)


# Adding preliminary information
data = document.add_paragraph('')
data.add_run(f'Programs offered 			:  ').bold = True
data.add_run(f'BTech CSE, MTech CSE, MTech AI\n')
data.add_run(f'No. of Students 			: B.Tech.     :  ').bold = True
data.add_run(f"{BTech_students_count(students)}\n")
data.add_run(f'					  M.Tech.    :  ').bold = True
data.add_run(f"{MTech_students_count(students)}\n")
data.add_run(f'					  MS (R )     :  ').bold = True
data.add_run(f"{MS_students_count(students)}\n")
data.add_run(f'					  PhD  	       :  ').bold = True
data.add_run(f"{PhD_students_count(students)}\n")
data.add_run(f'Head of the Department		:  ').bold = True
data.add_run(f"{Department['Hod']['username']} \n")
data.add_run(f'No. of faculty members		:  ').bold = True
data.add_run(f"{get_faculty_count(faculties)}\n")
data.add_run(f'No. of staff members		:\n').bold = True
data.add_run(f'					 Technical Staff :  ').bold = True
data.add_run(f"{get_tech_staff_count(staff)}\n")
data.add_run(f'					 Administrative Staff :  ').bold = True
data.add_run(f"{get_admin_staff_count(staff)}\n")
data.add_run(f'Thrust Area				:  ').bold = True
data.add_run("\n")
data.add_run(f'No. of Publications 			:  ').bold = True
data.add_run(f"{get_publications_count(publications_records)}\n")

# Adding Faculty details

document.add_paragraph('')
document.add_heading('Faculty Members', 1)
for i in faculties:
    para = document.add_paragraph()
    para.add_run("Name : ").bold = True
    para.add_run(f"{i['username']}\n")
    para.add_run("Designation : ").bold = True
    para.add_run(f"{i['designation']}\n")
    para.add_run("PhD Institute : ").bold = True
    para.add_run(f"{i['phd_institute']}\n")
    para.add_run("Areas : ").bold = True
    para.add_run(f"{i['areas']}\n")

# Adding Ongoing Activities
document.add_paragraph('')
document.add_heading('Ongoing Activities', 1)

# Adding Facilities
document.add_paragraph('')
document.add_heading('Facilities', 1)
facilities = document.add_paragraph('')
facilities.add_run(f'No. of Labs 					: 	UG	       :  ').bold = True
facilities.add_run(f'{get_ug_labs_counts(students)}\n')
facilities.add_run(f'							PG	       :  ').bold = True
facilities.add_run(f'{get_pg_labs_counts(students)}\n')
facilities.add_run(f'							Research   :  ').bold = True
facilities.add_run(f'{get_research_labs_counts(research_labs)}\n')


document.add_heading(f'AWARDS AND HONOURS {year}	(faculty)', 1)
document.add_paragraph('')
def add_table(document, headings, titles,records):
    # align = 'center'

    table = document.add_table(rows=1, cols=len(headings)+1, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Sr.No"
    for i in range(len(headings)):
        hdr_cells[i+1].text = titles[i]
    
    count = 1
    for record in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(count)
        for i in range(len(headings)):
            row_cells[i+1].text = record[headings[i]]
        count+=1


achievement_headings = ('faculty','title', 'type')
add_table(document, achievement_headings, achievement_headings,  achievement_records)
document.add_paragraph('')


document.add_heading('AWARDS AND HONOURS {year}	(student)', 1)
document.add_paragraph('')
student_achievement_headings = ('student','title', 'type')
add_table(document, student_achievement_headings, student_achievement_headings, student_achievement_records)
document.add_paragraph('')

document.add_heading('LECTURES BY VISITING EXPERTS', 1)
document.add_paragraph('')
event_headings = ('speakers', 'title', 'date')
event_titles = ('Name of the experts with affiliation', 'Topic', 'date')
add_table(document, event_headings, event_titles, event_records)
document.add_paragraph('')



document.add_heading('VISITS ABROAD BY THE FACULTY', 1)
document.add_paragraph('')
visit_headings = ('user', 'venue', 'title', 'from_date')
visit_titles = ('Name of the faculty member', 'Country', 'Visit Details', 'Date of visit')

add_table(document, visit_headings, visit_titles, visit_records)
document.add_paragraph('')


document.add_heading('VISITS ABROAD BY THE STUDENTS', 1)
document.add_paragraph('')
visit_headings = ('user', 'venue', 'title', 'from_date')
visit_titles = ('Name of the Student', 'Country', 'Visit Details', 'Date of visit')

add_table(document, visit_headings, visit_titles, visit_records)
document.add_paragraph('')


document.add_heading('5) Any other important information', 1)
document.add_paragraph('')
document.add_paragraph('')
document.add_paragraph('(Note: The data will be included from April 1, 2019 to March 31, 2020)')
document.add_page_break()

document.save('filename.docx')