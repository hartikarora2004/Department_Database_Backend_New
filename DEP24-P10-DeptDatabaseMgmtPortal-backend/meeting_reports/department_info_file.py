from docx import Document
from docx.shared import Inches
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt,RGBColor
from department.models import Department
from student_details.models import studentDetails
from faculty_details.models import facultyDetails
from staff_details.models import staffDetails
from usercustom.models import CustomUser
from research_lab.models import ResearchLab

def add_table(document, headings, titles,records):
    # align = 'center'

    table = document.add_table(rows=1, cols=len(headings)+1, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(1.4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Sr. No."
    for i in range(len(headings)):
        hdr_cells[i+1].text = titles[i]
    
    print("iterating records")
    count = 1
    for record in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(count)
        for i in range(len(headings)):
            value = getattr(record, headings[i])
            if(value== None):
                row_cells[i+1].text = "NA"
                continue
            if headings[i] == 'amount_invested':
                value = str(value/10000000) + " Cr"
            row_cells[i+1].text = str(value)
        count+=1
    print("table added")


def get_admin_staff_count(department):
    st = staffDetails.objects.filter(staff__is_current=True,type = 'Administrative Staff').filter(staff__department=department)
    return len(st)

def get_tech_staff_count(department):
    st = staffDetails.objects.filter(staff__is_current=True,type = 'Technical Staff').filter(staff__department=department)
    return len(st)

def BTech_students_count(batches):
    print(batches)
    print("a")
    x = studentDetails.objects.filter(batch__in=batches)
    print("b")
    det = x.filter(degree='B.Tech')
    print("c")
    count = len(det)
    return count

def MTech_students_count(batches):
    det = studentDetails.objects.filter(batch__in=batches).filter(degree='M.Tech')
    count = len(det)
    return count

def MS_students_count(batches):
    det = studentDetails.objects.filter(batch__in=batches).filter(degree='M.Sc')
    count = len(det)
    return count

def PhD_students_count(batches):
    det = studentDetails.objects.filter(batch__in=batches).filter(degree='Ph.D')
    count = len(det)
    return count

def get_faculty_count(facs):
    return facs.count()

def get_publications_count(publications_records):
    return publications_records.count()

def get_ug_labs_counts(labs):
    return labs.filter(lab_type='UG Lab').count()

def get_total_labs_counts(labs):
    return len(labs)

def get_pg_labs_counts(labs):
    return len(labs.filter(lab_type='PG Lab'))

def get_research_labs_counts(labs):
    return len(labs.filter(lab_type='Research Lab'))

def generate_file_now(filename, batches, Department,start_date, end_date, achievement_records,student_achievement_records,event_records,visit_records,st_visit_records, publications_records, project_records, year):
    print("generating file")
    print("start: ",start_date)
    print("end: ",end_date)
    document = Document('static/editfile.docx')
    section = document.sections[0]
    header = section.header
    # print("dates: ", start_date, end_date)
    if start_date!="" and end_date!="":
        header.paragraphs[0].add_run(f'Annual Report {start_date} to {end_date}\n')
    else:
        header.paragraphs[0].add_run(f'Annual Report {year}-{year+1}\n')
        

    # batches = [1,2,3,4,5]
    # Adding Title
    document.add_heading(f'DEPARTMENT OF {Department.name}'.upper(), 0)
    print("stage 1")
    # Adding preliminary information
    facs = facultyDetails.objects.filter(faculty__user_type='fc').filter(faculty__is_current=True).filter(faculty__department=Department).order_by('faculty__first_name')
    print(len(facs))
    # # # print(get_faculty_count(facs))
    # Check for matching records
    matching_faculty = CustomUser.objects.filter(user_type='fc',is_current=True, department_id=Department)
    matching_staff = CustomUser.objects.filter(user_type='st',is_current=True, department_id=Department)
    print('1.0')
    data = document.add_paragraph('')
    data.add_run(f'Programs offered 			:  ').bold = True
    data.add_run(f'{Department.programs_offered}\n')
    print("1.01")
    data.add_run(f'No. of Students 			: B.Tech.     :  ').bold = True
    data.add_run(f"{BTech_students_count(batches)}\n")
    data.add_run(f'					  M.Tech.    :  ').bold = True
    data.add_run(f"{MTech_students_count(batches)}\n")
    print("1.1")
    data.add_run(f'					  MS (R )     :  ').bold = True
    data.add_run(f"{MS_students_count(batches)}\n")
    data.add_run(f'					  PhD  	       :  ').bold = True
    data.add_run(f"{PhD_students_count(batches)}\n")
    data.add_run(f'Head of the Department		:  ').bold = True
    # data.add_run(f"{Department.Hod.username} \n")
    data.add_run(f"{Department.Hod.__str__()} \n")
    data.add_run(f'No. of Faculty Members		:  ').bold = True
    data.add_run(f"{matching_faculty.count()}\n")
    print("1.2")
    data.add_run(f'No. of Staff Members			:  ').bold = True
    data.add_run(f"{matching_staff.count()}\n")
    data.add_run(f'					  Technical Staff :  ').bold = True
    data.add_run(f"{get_tech_staff_count(Department)}\n")
    data.add_run(f'					  Administrative Staff :  ').bold = True
    data.add_run(f"{get_admin_staff_count(Department)}\n")
    print("1.3")
    data.add_run(f'No. of Publications 			:  ').bold = True
    data.add_run(f"{get_publications_count(publications_records)}\n")
    # print("stage 2")
    # Adding Faculty details
    document.add_paragraph('')
    document.add_heading('Faculty Members', 1)
    for i in facs:
        
        para = document.add_paragraph()
        para.add_run("Name : ").bold = True
        para.add_run(f"{str(i.faculty)}\n")
        para.add_run("Designation : ").bold = True
        para.add_run(f"{i.designation if i.designation else 'N.A.'}\n")
        para.add_run("PhD Institute : ").bold = True
        para.add_run(f"{i.phd_instuition if i.phd_instuition else 'N.A.'}\n")            
        para.add_run("Areas : ").bold = True
        para.add_run(f"{i.fields_of_interest if i.fields_of_interest else 'N.A.'}\n")   
    print("stage 3")
    # Adding Ongoing Activities
    # document.add_paragraph('')
    # document.add_heading('Ongoing Activities', 1)

    # Adding facilities
    labs = ResearchLab.objects.all()
    labs = labs.order_by('lab_type')
    # document.add_paragraph('')
    document.add_heading('Facilities', 1)
    facilities = document.add_paragraph('')
    facilities.add_run(f'No. of Labs 				: ').bold = True
    facilities.add_run(f'{get_total_labs_counts(labs)}\n')
    facilities.add_run(f'					UG:  ').bold = True
    facilities.add_run(f'{get_ug_labs_counts(labs)}\n')
    temp = get_pg_labs_counts(labs)
    # if temp == None and temp == 0:
    facilities.add_run(f'					PG:  ').bold = True
    facilities.add_run(f'{get_pg_labs_counts(labs)}\n')
    # temp = get_research_labs_counts(labs)
    # if temp == None and temp == 0:
    facilities.add_run(f'					Research:  ').bold = True
    facilities.add_run(f'{get_research_labs_counts(labs)}\n')
    if(labs!=None and len(labs)):
        for i in labs:
            para = document.add_paragraph()
            para.add_run(f"Name of the Lab				: ").bold = True
            para.add_run(f"{i.name}\n")
            para.add_run(f"Name of the Head of the Research Lab 	: ").bold = True
            para.add_run(f"{i.Head}\n")
            if i.equipments!=None and i.equipments !="":
                para.add_run(f"Name of the Equipments	  : ").bold = True
                para.add_run(f"{i.equipments}\n")


    # adding awards/honors
    if achievement_records!=None and len(achievement_records):
        document.add_heading(f'AWARDS AND HONORS {year}	(Faculty)', 1)
        document.add_paragraph('')
        achievement_headings = ('users','title', 'type')
        achievement_titles = ('Faculty','Title', 'Type')
        for achievement_record in achievement_records:
            value = getattr(achievement_record, 'type')
            expanded_value = 'NA'
            if(value == 'HC'): expanded_value = 'Hackathon'
            elif(value == 'CP'): expanded_value = 'Competition'
            elif(value == 'IN'): expanded_value = 'Internship'
            elif(value == 'O'): expanded_value = 'Other'
            setattr(achievement_record, 'type', expanded_value)
        add_table(document, achievement_headings, achievement_titles,  achievement_records)
        document.add_paragraph('')


    # adding awards/honors student
    if student_achievement_records!=None and len(student_achievement_records):
        document.add_heading(f'AWARDS AND HONORS {year}	(Student)', 1)
        document.add_paragraph('')
        student_achievement_headings = ('users','title', 'type')
        student_achievement_titles = ('Students','Title', 'Type')
        for student_achievement_record in student_achievement_records:
            value = getattr(student_achievement_record, 'type')
            expanded_value = 'NA'
            if(value == 'HC'): expanded_value = 'Hackathon'
            elif(value == 'CP'): expanded_value = 'Competition'
            elif(value == 'IN'): expanded_value = 'Internship'
            elif(value == 'O'): expanded_value = 'Other'
            setattr(student_achievement_record, 'type', expanded_value)
        add_table(document, student_achievement_headings, student_achievement_titles, student_achievement_records)
        document.add_paragraph('')
    print("stage 5")

    if event_records!=None and len(event_records):
        document.add_heading('LECTURES BY VISITING EXPERTS', 1)
        document.add_paragraph('')
        event_headings = ('speakers', 'title', 'date')
        event_titles = ('Name of the Expert(s) with Affiliation', 'Topic', 'Date(YYYY-MM-DD)')
        add_table(document, event_headings, event_titles, event_records)
        document.add_paragraph('')
    
    if visit_records!=None and len(visit_records):
        document.add_heading('VISITS ABROAD BY THE FACULTY', 1)
        document.add_paragraph('')
        visit_headings = ('users', 'venue', 'title', 'from_date')
        visit_titles = ('Name of the Faculty Member', 'Country', 'Visit Details', 'Date of Visit')
        add_table(document, visit_headings, visit_titles, visit_records)
        document.add_paragraph('')
    print("stage 6")
    if st_visit_records!=None and len(st_visit_records):
        document.add_heading('VISITS ABROAD BY THE STUDENTS', 1)
        document.add_paragraph('')
        visit_headings = ('users', 'venue', 'title', 'from_date')
        visit_titles = ('Name of the Student', 'Country', 'Visit Details', 'Date of Visit')
        add_table(document, visit_headings, visit_titles, st_visit_records)
        document.add_paragraph('')
    print("stage 7")
    if project_records!=None and len(project_records):
        document.add_heading('MAJOR RESEARCH PROJECTS (Ongoing/Completed)', 1)
        document.add_paragraph('')
        project_headings = ('users', 'title', 'start_date', 'status', 'investors', 'amount_invested')
        project_titles = ('Name of the Faculty Member', 'Title of the Project', 'Start Date', 'Ongoing/Completed', 'Name of the Funding Agency', 'Amount Invested (in Rs. Crores)')
        add_table(document, project_headings, project_titles, project_records)
        document.add_paragraph('')
    print("stage 8")


    document.add_heading('PUBLICATIONS', 1)
    print(len(facs))
    for i in facs:
        # print(i.faculty)
        temp = publications_records.filter(authors__in = [i.faculty])
        if temp!=None and len(temp):
            document.add_heading(f'Faculty : {i.faculty.pub_name()}', 2)
            pub_titles = document.add_paragraph('')
            pub_titles.add_run('Papers:\n').bold = True
            for j in temp:
                if j.description!=None and j.description!="":
                    pub_titles.add_run(f"{j.description}\n")
                else:
                    pub_titles.add_run(f"{j.title} by {j.authors_text} \n")
                    

    # document.add_heading('5) Any other Important Information', 1)
    # document.add_paragraph('\n\n\n(Note: The data will be included from April 1, 2019 to March 31, 2020)')
    document.add_page_break()

    document.save(filename)