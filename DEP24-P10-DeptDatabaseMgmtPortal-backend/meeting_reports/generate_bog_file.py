from docx import Document
from docx.shared import Inches
from docx.shared import Cm
from datetime import datetime
from docx.enum.table import WD_TABLE_ALIGNMENT


def add_table(document, headings, titles, records, column_widths):
    table = document.add_table(rows=1, cols=len(headings)+1, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = True
    for i in range(len(headings)):
        table.columns[i].width = Cm(column_widths[i])
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Sl. No."
    for i in range(len(headings)):
        hdr_cells[i+1].text = titles[i]
    print("iterating records")
    count = 1
    for record in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(count)
        for i in range(len(headings)):
            value = getattr(record, headings[i])
            if(value== None):
                row_cells[i+1].text = "NA"
                continue
            row_cells[i+1].text = str(value)
        count+=1
    print("table added")



def generate_file(filename, start_date, end_date, achievement_records, student_achievement_records, event_records, visit_records):
    print("generating file")
    document = Document()
    section = document.sections[0]
    header = section.header
    header.paragraphs[0].add_run(f'BoG Report {start_date} to {end_date}\n')
    counter  = 1
    document.add_heading('Information for BoG Meeting', 0)
    if achievement_records!=None and len(achievement_records):
        document.add_heading(f'{counter}) Award/Achievement by Faculty Members', 1)
        achievement_headings = ('users','title', 'type')
        achievement_titles = ('Faculty','Title', 'Type')
        for achievement_record in achievement_records:
            value = getattr(achievement_record, 'type')
            expanded_value = 'NA'
            if(value == 'HC'): expanded_value = 'Hackathon'
            elif(value == 'CP'): expanded_value = 'Competition'
            elif(value == 'IN'): expanded_value = 'Internship'
            elif(value == 'O'): expanded_value = 'Other'
            setattr(achievement_record, 'type', expanded_value)
        achievement_column_widths = [1.4, 4, 6.2, 3.5]
        print("Records", achievement_records)
        add_table(document, achievement_headings, achievement_titles,  achievement_records, achievement_column_widths)
        counter+=1
    print("added achievement table")
    if student_achievement_records!=None and len(student_achievement_records):
        document.add_heading(f'{counter}) Award/Achievement by Students', 1)
        student_achievement_headings = ('users','title', 'type')
        student_achievement_titles = ('Student','Title', 'Type')
        for student_achievement_record in student_achievement_records:
            value = getattr(student_achievement_record, 'type')
            expanded_value = 'NA'
            if(value == 'HC'): expanded_value = 'Hackathon'
            elif(value == 'CP'): expanded_value = 'Competition'
            elif(value == 'IN'): expanded_value = 'Internship'
            elif(value == 'O'): expanded_value = 'Other'
            setattr(student_achievement_record, 'type', expanded_value)
        student_achievement_column_widths = [1.4, 4, 6.2, 3.5]
        add_table(document, student_achievement_headings, student_achievement_titles, student_achievement_records, student_achievement_column_widths)
        counter+=1

    print("added student achievement table")
    if event_records!=None and len(event_records):
        document.add_heading(f'{counter}) Workshop/Conferences/Seminars organised by Department ', 1)
        event_headings = ('speakers', 'title', 'date', 'venue', 'number_of_participants', 'description')
        event_titles = ('Name of Speaker(s)/ Chief Guest(s)', 'Title', 'Date', 'Venue', 'No. of Participants', 'Any Other Important Information')
        event_column_widths = [1.4, 2.7, 2.7, 2.4, 1.6, 1.6, 2.4]
        add_table(document, event_headings, event_titles, event_records, event_column_widths)
        counter+=1

    print("added event table")
    if visit_records!=None and len(visit_records):
        document.add_heading(f'{counter}) Faculty Participation as Invited Speaker / Session Chair', 1)
        visit_headings = ('users', 'title', 'from_date', 'venue', 'description')
        visit_titles = ('Name of Faculty', 'Title of the Event', 'Date of Visit', 'Venue', 'Any Other Important Information')
        visit_column_widths = [1.4, 3, 3, 2.4, 2.4, 2.6]
        add_table(document, visit_headings, visit_titles, visit_records, visit_column_widths)
        counter+=1

    print("added visit table")
    # document.add_heading(f'{counter}) Any Other Important Information', 1)

    document.add_page_break()

    document.save(filename)